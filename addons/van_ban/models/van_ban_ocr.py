# -*- coding: utf-8 -*-

import base64
import hashlib
import io
import zipfile
import os
import shutil

import json
import logging

from odoo import api, fields, models, _
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

try:
    from PIL import Image
except Exception:
    Image = None

try:
    from PIL import ImageOps, ImageFilter
except Exception:
    ImageOps = None
    ImageFilter = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    import requests
except Exception:
    requests = None

from .ocr_utils import ocr_image_bytes, fix_spacing_artifacts


class VanBanOCR(models.Model):
    _name = 'van_ban_ocr'
    _description = 'OCR - Trích xuất nội dung từ file'
    _inherit = ['mail.thread', 'mail.activity.mixin']

    name = fields.Char('Tên OCR', required=True, tracking=True)
    file_dinh_kem = fields.Binary('File', attachment=True, tracking=True)
    ten_file = fields.Char('Tên file', tracking=True)

    history_ids = fields.One2many('van_ban_ocr_history', 'ocr_id', string='Lịch sử', readonly=True)

    loai_file = fields.Selection(
        [('docx', 'DOCX'), ('image', 'Ảnh'), ('unknown', 'Không xác định')],
        string='Loại file',
        compute='_compute_loai_file',
        store=True,
        readonly=True,
    )

    noi_dung_trich_xuat = fields.Text('Nội dung trích xuất', tracking=True)
    loi_xu_ly = fields.Text('Lỗi xử lý', readonly=True)
    
    # AI Summary
    ai_summary = fields.Text('Tóm tắt AI', readonly=True, tracking=True,
                             help='Tóm tắt nội dung văn bản bởi AI')
    ai_summary_at = fields.Datetime('Tóm tắt lúc', readonly=True)

    def _get_ocr_provider(self):
        return self.env['ir.config_parameter'].sudo().get_param('van_ban.ocr_provider') or 'local'

    def _get_ocrspace_settings(self):
        ICP = self.env['ir.config_parameter'].sudo()
        return {
            'api_key': (ICP.get_param('van_ban.ocrspace_api_key') or '').strip(),
            'language': (ICP.get_param('van_ban.ocrspace_language') or 'vie').strip(),
            'engine': (ICP.get_param('van_ban.ocrspace_engine') or '2').strip(),
        }

    def _get_docx_ocr_settings(self):
        ICP = self.env['ir.config_parameter'].sudo()
        enabled = ICP.get_param('van_ban.ocr_docx_images')
        # get_param returns strings; treat '', '0', 'false' as False
        enabled = str(enabled).lower() not in ('', '0', 'false', 'none')
        try:
            max_images = int(ICP.get_param('van_ban.ocr_docx_max_images') or 10)
        except Exception:
            max_images = 10
        if max_images <= 0:
            max_images = 10
        return {'enabled': enabled, 'max_images': max_images}

    def _extract_docx_images(self, docx_bytes, *, max_images=10):
        """Return list of (filename, bytes) for images embedded in a DOCX."""
        images = []
        try:
            with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zf:
                for name in zf.namelist():
                    # Images are typically stored under word/media/
                    if not name.startswith('word/media/'):
                        continue
                    if len(images) >= max_images:
                        break
                    try:
                        raw = zf.read(name)
                    except Exception:
                        continue
                    if raw:
                        images.append((os.path.basename(name), raw))
        except Exception:
            return []
        return images

    def action_clear_file(self):
        for record in self:
            record.write({'file_dinh_kem': False, 'ten_file': False})

    @api.onchange('ten_file')
    def _onchange_ten_file_suggest_name(self):
        for record in self:
            if record.ten_file and (not record.name or record.name == _('OCR')):
                base_name = os.path.splitext(record.ten_file)[0]
                record.name = base_name or _('OCR')

    @api.depends('file_dinh_kem', 'ten_file')
    def _compute_loai_file(self):
        for record in self:
            record.loai_file = record._detect_file_type() if record.file_dinh_kem else 'unknown'

    def _detect_file_type(self):
        self.ensure_one()
        file_data = base64.b64decode(self.file_dinh_kem) if self.file_dinh_kem else b''
        file_name = (self.ten_file or '').lower()
        header = file_data[:16] if file_data else b''

        is_docx = file_name.endswith('.docx') or header.startswith(b'PK')
        is_image = file_name.endswith(('.png', '.jpg', '.jpeg')) or (
            header.startswith(b'\x89PNG\r\n\x1a\n') or
            header.startswith(b'\xff\xd8\xff') or
            header.startswith(b'GIF87a') or
            header.startswith(b'GIF89a') or
            header[:2] == b'BM'
        )

        if is_docx:
            return 'docx'
        if is_image:
            return 'image'
        return 'unknown'

    def _compute_sha1(self, raw_bytes):
        if not raw_bytes:
            return False
        return hashlib.sha1(raw_bytes).hexdigest()

    def _ensure_ocr_dependencies(self):
        missing_python = []
        if pytesseract is None:
            missing_python.append('pytesseract')
        if Image is None:
            missing_python.append('Pillow')

        if missing_python:
            raise UserError(_(
                'Thiếu thư viện OCR (%s). Cài bằng: pip install -r addons/van_ban/requirements.txt'
            ) % ', '.join(missing_python))

        if not shutil.which('tesseract'):
            raise UserError(_(
                'Thiếu chương trình "tesseract" trên hệ điều hành. Trên Ubuntu/Debian: '
                'sudo apt-get install -y tesseract-ocr tesseract-ocr-vie'
            ))

    def _ensure_ocr_api_dependencies(self):
        if requests is None:
            raise UserError(_(
                'Thiếu thư viện gọi API (requests). Cài bằng: pip install requests'
            ))

    def _ocr_via_ocrspace(self, image_bytes):
        """Call OCR.Space API. Returns extracted text."""
        self.ensure_one()
        self._ensure_ocr_api_dependencies()

        settings = self._get_ocrspace_settings()
        if not settings['api_key']:
            raise UserError(_('Chưa cấu hình OCR.Space API Key trong Cài đặt.'))

        # OCR.Space endpoint
        url = 'https://api.ocr.space/parse/image'

        # OCR.Space expects multipart form: file + params
        files = {
            'file': ('upload.png', image_bytes),
        }
        data = {
            'apikey': settings['api_key'],
            'language': settings['language'] or 'vie',
            'OCREngine': settings['engine'] or '2',
            'isOverlayRequired': 'false',
            # keep it conservative; user can improve quality by preprocessing before upload if needed
        }

        try:
            resp = requests.post(url, files=files, data=data, timeout=60)
        except Exception as e:
            raise UserError(_('Không gọi được OCR API: %s') % str(e))

        if resp.status_code != 200:
            raise UserError(_('OCR API trả về lỗi HTTP %s: %s') % (resp.status_code, (resp.text or '')[:300]))

        try:
            payload = resp.json()
        except Exception:
            # Some proxies may mangle JSON content-type
            try:
                payload = json.loads(resp.text or '{}')
            except Exception:
                raise UserError(_('OCR API trả về dữ liệu không hợp lệ.'))

        if payload.get('IsErroredOnProcessing'):
            error_messages = payload.get('ErrorMessage')
            if isinstance(error_messages, list):
                error_messages = '; '.join([m for m in error_messages if m])
            raise UserError(_('OCR API báo lỗi: %s') % (error_messages or 'Unknown error'))

        parsed = payload.get('ParsedResults') or []
        text_parts = []
        for item in parsed:
            t = (item.get('ParsedText') or '').strip()
            if t:
                text_parts.append(t)
        return ('\n'.join(text_parts)).strip()

    def _ensure_docx_dependencies(self):
        if Document is None:
            raise UserError(_(
                'Thiếu thư viện đọc DOCX (python-docx). Cài bằng: pip install -r addons/van_ban/requirements.txt'
            ))

    def action_trich_xuat(self):
        """Trích xuất văn bản từ file, tự động lưu không cần ấn Save."""
        self.ensure_one()
        if not self.file_dinh_kem:
            raise UserError(_('Vui lòng upload file trước khi trích xuất.'))
        self._run_extract()
        
        # Return notification with reload to refresh view
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': _('Trích xuất hoàn thành'),
                'message': _('Đã trích xuất văn bản thành công.'),
                'type': 'success',
                'sticky': False,
                'next': {'type': 'ir.actions.client', 'tag': 'reload'},
            }
        }

    def _run_extract(self):
        self.ensure_one()

        _logger.info("Starting OCR extraction for file: %s", self.ten_file)

        self.loi_xu_ly = False
        self.noi_dung_trich_xuat = False

        loai = self._detect_file_type()
        _logger.info("Detected file type: %s", loai)
        file_data = base64.b64decode(self.file_dinh_kem) if self.file_dinh_kem else b''
        checksum = self._compute_sha1(file_data)
        lang_used = False
        tesseract_config = False
        provider_used = False

        try:
            if loai == 'docx':
                self._ensure_docx_dependencies()
                doc = Document(io.BytesIO(file_data))
                parts = []

                provider = self._get_ocr_provider()
                provider_used = provider

                for p in doc.paragraphs:
                    text = (p.text or '').strip()
                    if text:
                        parts.append(text)

                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = (cell.text or '').strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            parts.append(' | '.join(row_text))

                # Optional: OCR embedded images inside DOCX (useful for scanned DOCX)
                docx_cfg = self._get_docx_ocr_settings()
                if docx_cfg.get('enabled'):
                    images = self._extract_docx_images(file_data, max_images=docx_cfg.get('max_images', 10))
                    if images:
                        ocr_texts = []
                        for (img_name, img_bytes) in images:
                            try:
                                if provider == 'ocrspace':
                                    lang_used = self._get_ocrspace_settings().get('language') or 'vie'
                                    txt = self._ocr_via_ocrspace(img_bytes)
                                else:
                                    self._ensure_ocr_dependencies()
                                    if ImageOps is None or ImageFilter is None:
                                        raise UserError(_('Thiếu thư viện Pillow để xử lý ảnh OCR. Cài bằng: pip install -r addons/van_ban/requirements.txt'))
                                    lang_used = 'vie+eng'
                                    tesseract_config = '--oem 3 --psm 6 -c preserve_interword_spaces=1'
                                    txt = ocr_image_bytes(
                                        Image,
                                        ImageOps,
                                        ImageFilter,
                                        pytesseract,
                                        img_bytes,
                                        lang=lang_used,
                                        config=tesseract_config,
                                    )
                                txt = (txt or '').strip()
                                if txt:
                                    ocr_texts.append(_('--- OCR từ ảnh trong DOCX: %s ---\n%s') % (img_name, txt))
                            except Exception:
                                # Ignore per-image OCR failures to avoid blocking DOCX extraction
                                continue

                        if ocr_texts:
                            parts.extend(ocr_texts)

                self.noi_dung_trich_xuat = fix_spacing_artifacts(('\n'.join(parts)).strip()) or False

            elif loai == 'image':
                provider = self._get_ocr_provider()
                provider_used = provider

                if provider == 'ocrspace':
                    # OCR via external API
                    lang_used = self._get_ocrspace_settings().get('language') or 'vie'
                    text = self._ocr_via_ocrspace(file_data)
                    self.noi_dung_trich_xuat = fix_spacing_artifacts((text or '').strip()) or False
                else:
                    # Local OCR (Tesseract)
                    self._ensure_ocr_dependencies()
                    if ImageOps is None or ImageFilter is None:
                        raise UserError(_('Thiếu thư viện Pillow để xử lý ảnh OCR. Cài bằng: pip install -r addons/van_ban/requirements.txt'))
                    lang_used = 'vie+eng'
                    tesseract_config = '--oem 3 --psm 6 -c preserve_interword_spaces=1'
                    text = ocr_image_bytes(
                        Image,
                        ImageOps,
                        ImageFilter,
                        pytesseract,
                        file_data,
                        lang=lang_used,
                        config=tesseract_config,
                    )
                    self.noi_dung_trich_xuat = fix_spacing_artifacts((text or '').strip()) or False

            else:
                raise UserError(_('Chỉ hỗ trợ .docx, .png, .jpg/.jpeg.'))

        except UserError:
            raise
        except Exception as e:
            _logger.error("OCR extraction failed: %s", str(e), exc_info=True)
            self.loi_xu_ly = str(e)
            self.noi_dung_trich_xuat = False

        self.env['van_ban_ocr_history'].sudo().create({
            'ocr_id': self.id,
            'name': _('Trích xuất %s') % (fields.Datetime.now() or ''),
            'user_id': self.env.user.id,
            'file_dinh_kem': self.file_dinh_kem,
            'ten_file': self.ten_file,
            'loai_file': loai,
            'checksum_sha1': checksum,
            'lang': lang_used,
            'tesseract_config': tesseract_config,
            'noi_dung_trich_xuat': self.noi_dung_trich_xuat,
            'loi_xu_ly': self.loi_xu_ly,
        })

    def action_ai_summarize(self):
        """Tóm tắt nội dung trích xuất bằng AI."""
        self.ensure_one()
        
        if not self.noi_dung_trich_xuat:
            raise UserError(_('Chưa có nội dung trích xuất. Vui lòng trích xuất văn bản trước.'))
        
        ai_service = self.env['ai.service']
        if not ai_service.is_available():
            raise UserError(_('AI Service chưa được cấu hình. Vui lòng vào Settings > AI Integration.'))
        
        text = self.noi_dung_trich_xuat
        
        if len(text) < 50:
            raise UserError(_('Nội dung quá ngắn để tóm tắt.'))
        
        try:
            summary = ai_service.summarize_text(
                text=text,
                max_words=300,
                focus='các điểm chính, thông tin quan trọng, bên liên quan và nội dung cốt lõi của văn bản',
                model_name=self._name,
                record_id=self.id
            )
            
            self.write({
                'ai_summary': summary,
                'ai_summary_at': fields.Datetime.now(),
            })
            
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'title': _('Tóm tắt hoàn thành'),
                    'message': _('AI đã tóm tắt văn bản thành công.'),
                    'type': 'success',
                    'sticky': False,
                    'next': {'type': 'ir.actions.client', 'tag': 'reload'},
                }
            }
        except Exception as e:
            _logger.error('AI Summary failed: %s', str(e), exc_info=True)
            raise UserError(_('Lỗi khi tóm tắt: %s') % str(e))

    @api.model_create_multi
    def create(self, vals_list):
        records = super(VanBanOCR, self).create(vals_list)
        for record in records:
            if record.file_dinh_kem:
                record._run_extract()
        return records

    def write(self, vals):
        res = super(VanBanOCR, self).write(vals)
        if 'file_dinh_kem' in vals or 'ten_file' in vals:
            for record in self:
                if record.file_dinh_kem:
                    record._run_extract()
                else:
                    record.noi_dung_trich_xuat = False
                    record.loi_xu_ly = False
        return res
